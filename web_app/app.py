"""
Financial Planner Web Application
Flask backend for the Financial Planning Agent System
"""
import os
import json
import uuid
import re
from datetime import datetime
from io import BytesIO
from flask import Flask, render_template, request, jsonify, session, send_file, make_response
from flask_cors import CORS
from flask_session import Session
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
from reportlab.lib.colors import HexColor, black, white
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from config import get_config
from agents import AgentState, OrchestratorAgent
from visualizations import get_visualizations
from langchain_openai import ChatOpenAI
from langchain_core.messages import HumanMessage


# Initialize Flask app
def create_app():
    app = Flask(__name__, template_folder='templates', static_folder='static')
    
    # Load configuration
    config = get_config()
    app.config.from_object(config)
    
    # Initialize extensions
    CORS(app)
    Session(app)
    
    return app


app = create_app()

# Global session storage for planning results
planning_sessions = {}


# ============================================================================
# API ENDPOINTS
# ============================================================================

@app.route('/')
def index():
    """Landing page with agent/tool selection"""
    return render_template('index.html')


@app.route('/api/plans', methods=['GET'])
def get_available_plans():
    """Get available planning options"""
    plans = [
        {
            "id": "retirement",
            "name": "Retirement Planning",
            "description": "Calculate retirement needs, savings goals, and withdrawal strategies",
            "icon": "🏖️"
        },
        {
            "id": "insurance",
            "name": "Insurance Planning",
            "description": "Determine optimal insurance coverage for life, disability, and liability",
            "icon": "🛡️"
        },
        {
            "id": "estate",
            "name": "Estate Planning",
            "description": "Plan for wealth transfer, education funding, and tax minimization",
            "icon": "📋"
        },
        {
            "id": "wealth",
            "name": "Personal Wealth Management",
            "description": "Build comprehensive investment strategy and asset allocation",
            "icon": "💰"
        }
    ]
    return jsonify(plans)


@app.route('/api/planning/start', methods=['POST'])
def start_planning():
    """Start financial planning session"""
    try:
        data = request.json
        selected_plans = data.get('selected_plans', [])
        user_info = data.get('user_info', {})
        
        if not selected_plans:
            return jsonify({'error': 'No plans selected'}), 400
        
        if not user_info:
            return jsonify({'error': 'User information required'}), 400
        
        # Generate session ID
        session_id = str(uuid.uuid4())
        
        # Map plan IDs to names
        plan_mapping = {
            'retirement': 'Retirement Planning',
            'insurance': 'Insurance Planning',
            'estate': 'Estate Planning',
            'wealth': 'Personal Wealth Management'
        }
        
        selected_plan_names = [plan_mapping.get(p) for p in selected_plans if p in plan_mapping]
        
        # Initialize agent state
        initial_state = AgentState(
            messages=[],
            user_info=user_info,
            selected_plans=selected_plan_names,
            plan_summaries={},
            next_agent=""
        )
        
        # Initialize LLM
        llm = ChatOpenAI(
            model=app.config['OPENAI_MODEL'],
            temperature=app.config['OPENAI_TEMPERATURE'],
            api_key=app.config['OPENAI_API_KEY']
        )
        
        # Run orchestrator
        orchestrator = OrchestratorAgent(llm)
        final_state = orchestrator.route(initial_state)
        
        # Generate visualizations
        visualizations = get_visualizations(user_info, selected_plan_names)
        
        # Store session
        planning_sessions[session_id] = {
            'user_info': user_info,
            'selected_plans': selected_plan_names,
            'plan_summaries': final_state['plan_summaries'],
            'visualizations': visualizations,
            'conversation_history': [],
            'created_at': datetime.now().isoformat()
        }
        
        return jsonify({
            'session_id': session_id,
            'plan_summaries': final_state['plan_summaries'],
            'visualizations': visualizations,
            'status': 'success'
        })
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/planning/<session_id>', methods=['GET'])
def get_planning_results(session_id):
    """Get planning results for a session"""
    if session_id not in planning_sessions:
        return jsonify({'error': 'Session not found'}), 404
    
    session_data = planning_sessions[session_id]
    
    return jsonify({
        'user_info': session_data['user_info'],
        'selected_plans': session_data['selected_plans'],
        'plan_summaries': session_data['plan_summaries'],
        'visualizations': session_data['visualizations']
    })


@app.route('/api/chat/<session_id>', methods=['POST'])
def chat(session_id):
    """Chat endpoint for follow-up questions"""
    try:
        if session_id not in planning_sessions:
            return jsonify({'error': 'Session not found'}), 404
        
        data = request.json
        user_message = data.get('message', '')
        
        if not user_message:
            return jsonify({'error': 'Message required'}), 400
        
        session_data = planning_sessions[session_id]
        
        # Initialize LLM
        llm = ChatOpenAI(
            model=app.config['OPENAI_MODEL'],
            temperature=app.config['OPENAI_TEMPERATURE'],
            api_key=app.config['OPENAI_API_KEY']
        )
        
        # Build conversation context
        user_info = session_data['user_info']
        plan_summaries = session_data['plan_summaries']
        
        context_prompt = f"""You are a financial advisor assistant. You have access to the following information
about the client and their financial plan:

User Information:
{json.dumps(user_info, indent=2)}

Financial Plan Summaries:
{chr(10).join([f"### {plan}:{chr(10)}{summary}{chr(10)}" for plan, summary in plan_summaries.items()])}

Answer the user's questions based on this context. Be helpful, specific, and reference the plans when relevant.
Keep responses concise but informative."""
        
        # Build messages
        messages = [
            {"role": "system", "content": context_prompt}
        ]
        
        # Add conversation history
        for msg in session_data['conversation_history'][-8:]:  # Keep last 8 messages
            messages.append(msg)
        
        # Add current message
        messages.append({"role": "user", "content": user_message})
        
        # Get response
        response = llm.invoke(messages)
        assistant_message = response.content
        
        # Store in conversation history
        session_data['conversation_history'].append({"role": "user", "content": user_message})
        session_data['conversation_history'].append({"role": "assistant", "content": assistant_message})
        
        return jsonify({
            'message': assistant_message,
            'status': 'success'
        })
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/export/<session_id>', methods=['GET'])
def export_plan(session_id):
    """Export planning summary as JSON"""
    if session_id not in planning_sessions:
        return jsonify({'error': 'Session not found'}), 404
    
    session_data = planning_sessions[session_id]
    
    export_data = {
        'generated_at': datetime.now().isoformat(),
        'user_info': session_data['user_info'],
        'selected_plans': session_data['selected_plans'],
        'plan_summaries': session_data['plan_summaries']
    }
    
    return jsonify(export_data)


def clean_text_for_export(text):
    """Clean text for export by removing markdown and special characters"""
    if not text:
        return ""
    
    # Remove markdown headers
    text = re.sub(r'^#+\s*', '', text, flags=re.MULTILINE)
    # Remove bold markdown
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
    # Remove LaTeX escape characters
    text = re.sub(r'\\?\\\(', '(', text)
    text = re.sub(r'\\?\\\)', ')', text)
    # Clean extra whitespace
    text = re.sub(r'\n\s*\n', '\n\n', text)
    
    return text.strip()


def format_section_content(section_lines, content_style, bullet_style, important_style):
    """Helper function to format section content with proper styling"""
    formatted_content = []
    
    for line in section_lines:
        if not line.strip():
            continue
            
        # Check for bullet points
        if line.strip().startswith(('•', '-', '*')):
            bullet_text = line.strip()[1:].strip()
            formatted_content.append(Paragraph(f"• {bullet_text}", bullet_style))
        # Check for important text (contains keywords like "Important", "Note", "Warning")
        elif any(keyword in line.lower() for keyword in ['important', 'note', 'warning', 'caution', 'remember']):
            formatted_content.append(Paragraph(line, important_style))
        # Check for numerical data (currency, percentages)
        elif any(char in line for char in ['$', '%']) or re.search(r'\d+', line):
            # Create a table for financial data if it contains structured info
            if ':' in line and any(char in line for char in ['$', '%']):
                parts = line.split(':')
                if len(parts) == 2:
                    data = [[parts[0].strip(), parts[1].strip()]]
                    table = Table(data, colWidths=[3*inch, 3*inch])
                    table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, -1), HexColor('#f8f9fa')),
                        ('TEXTCOLOR', (0, 0), (-1, -1), black),
                        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
                        ('FONTSIZE', (0, 0), (-1, -1), 11),
                        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                        ('TOPPADDING', (0, 0), (-1, -1), 6),
                        ('GRID', (0, 0), (-1, -1), 0.5, HexColor('#e2e8f0'))
                    ]))
                    formatted_content.append(table)
                    formatted_content.append(Spacer(1, 6))
                else:
                    formatted_content.append(Paragraph(line, content_style))
            else:
                formatted_content.append(Paragraph(line, content_style))
        else:
            formatted_content.append(Paragraph(line, content_style))
    
    return formatted_content


@app.route('/api/export/<session_id>/pdf', methods=['GET'])
def export_plan_pdf(session_id):
    """Export planning summary as enhanced PDF with professional formatting"""
    if session_id not in planning_sessions:
        return jsonify({'error': 'Session not found'}), 404
    
    session_data = planning_sessions[session_id]
    
    # Create PDF buffer
    buffer = BytesIO()
    
    # Create PDF document with better margins
    doc = SimpleDocTemplate(buffer, pagesize=letter, 
                          rightMargin=0.75*inch, leftMargin=0.75*inch, 
                          topMargin=0.75*inch, bottomMargin=0.75*inch)
    
    # Get base styles
    styles = getSampleStyleSheet()
    
    # Enhanced custom styles
    title_style = ParagraphStyle(
        'EnhancedTitle',
        parent=styles['Heading1'],
        fontSize=26,
        spaceAfter=30,
        spaceBefore=0,
        textColor=HexColor('#1a365d'),
        fontName='Helvetica-Bold',
        alignment=1,  # Center
        borderWidth=2,
        borderColor=HexColor('#3182ce'),
        borderPadding=12,
        backColor=HexColor('#ebf8ff')
    )
    
    subtitle_style = ParagraphStyle(
        'EnhancedSubtitle',
        parent=styles['Normal'],
        fontSize=12,
        spaceAfter=25,
        textColor=HexColor('#4a5568'),
        fontName='Helvetica',
        alignment=1
    )
    
    section_heading_style = ParagraphStyle(
        'EnhancedSectionHeading',
        parent=styles['Heading2'],
        fontSize=18,
        spaceAfter=15,
        spaceBefore=25,
        textColor=HexColor('#2d3748'),
        fontName='Helvetica-Bold',
        backColor=HexColor('#f7fafc'),
        borderWidth=0.5,
        borderColor=HexColor('#e2e8f0'),
        borderPadding=8,
        leftIndent=0
    )
    
    subsection_style = ParagraphStyle(
        'EnhancedSubsection',
        parent=styles['Heading3'],
        fontSize=14,
        spaceAfter=12,
        spaceBefore=18,
        textColor=HexColor('#3182ce'),
        fontName='Helvetica-Bold'
    )
    
    content_style = ParagraphStyle(
        'EnhancedContent',
        parent=styles['Normal'],
        fontSize=11,
        spaceAfter=8,
        leading=16,
        textColor=HexColor('#2d3748'),
        fontName='Helvetica',
        alignment=0
    )
    
    bullet_style = ParagraphStyle(
        'EnhancedBullet',
        parent=styles['Normal'],
        fontSize=11,
        spaceAfter=6,
        leading=16,
        textColor=HexColor('#2d3748'),
        fontName='Helvetica',
        leftIndent=20,
        bulletIndent=10
    )
    
    important_style = ParagraphStyle(
        'EnhancedImportant',
        parent=styles['Normal'],
        fontSize=11,
        spaceAfter=12,
        leading=16,
        textColor=HexColor('#1a365d'),
        fontName='Helvetica-Bold',
        backColor=HexColor('#ebf8ff'),
        borderWidth=1,
        borderColor=HexColor('#3182ce'),
        borderPadding=8
    )
    
    # Build document content
    content = []
    
    # Header section
    content.append(Paragraph("Financial Planning Report", title_style))
    content.append(Paragraph(f"Generated on {datetime.now().strftime('%B %d, %Y')}", subtitle_style))
    
    # Add separator line
    line_data = [[''], ['']]
    line_table = Table(line_data, colWidths=[7*inch])
    line_table.setStyle(TableStyle([
        ('LINEBELOW', (0, 0), (-1, 0), 2, HexColor('#3182ce')),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
    ]))
    content.append(line_table)
    content.append(Spacer(1, 25))
    
    # Client Information Table
    content.append(Paragraph("Client Information", section_heading_style))
    user_info = session_data['user_info']
    
    # Create a formatted table for user info
    user_data = [
        ['Age:', str(user_info.get('age', 'N/A'))],
        ['Annual Income:', f"${user_info.get('annual_income', 0):,.2f}"],
        ['Current Savings:', f"${user_info.get('savings', 0):,.2f}"],
        ['Goals:', ', '.join(session_data.get('selected_plans', []))]
    ]
    
    user_table = Table(user_data, colWidths=[2*inch, 4*inch])
    user_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), HexColor('#f8f9fa')),
        ('BACKGROUND', (1, 0), (1, -1), white),
        ('TEXTCOLOR', (0, 0), (-1, -1), black),
        ('ALIGN', (0, 0), (0, -1), 'RIGHT'),
        ('ALIGN', (1, 0), (1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTNAME', (1, 0), (1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 11),
        ('ROWBACKGROUNDS', (0, 0), (-1, -1), [HexColor('#f8f9fa'), white]),
        ('GRID', (0, 0), (-1, -1), 0.5, HexColor('#e2e8f0')),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('TOPPADDING', (0, 0), (-1, -1), 8),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8)
    ]))
    content.append(user_table)
    content.append(Spacer(1, 25))
    
    # Executive Summary (if exists)
    if 'Executive Summary' in session_data['plan_summaries']:
        content.append(Paragraph("Executive Summary", section_heading_style))
        exec_summary = clean_text_for_export(session_data['plan_summaries']['Executive Summary'])
        
        # Format executive summary content
        summary_lines = exec_summary.split('\n')
        summary_content = format_section_content(summary_lines, content_style, bullet_style, important_style)
        content.extend(summary_content)
        content.append(PageBreak())
    
    # Individual plans with enhanced formatting
    for plan_name, summary in session_data['plan_summaries'].items():
        if plan_name != 'Executive Summary':
            content.append(Paragraph(plan_name, section_heading_style))
            
            clean_summary = clean_text_for_export(summary)
            
            # Parse and format the content
            lines = clean_summary.split('\n')
            plan_content = format_section_content(lines, content_style, bullet_style, important_style)
            content.extend(plan_content)
            
            content.append(Spacer(1, 25))
    
    # Add footer
    content.append(Spacer(1, 30))
    content.append(line_table)
    footer_style = ParagraphStyle(
        'Footer',
        parent=styles['Normal'],
        fontSize=9,
        textColor=HexColor('#718096'),
        alignment=1,
        spaceAfter=0
    )
    content.append(Paragraph("This financial plan is generated by AI and should be reviewed with a qualified financial advisor.", footer_style))
    content.append(Paragraph("For questions or updates, please consult your financial planning professional.", footer_style))
    
    # Build PDF
    doc.build(content)
    buffer.seek(0)
    
    # Create response
    response = make_response(buffer.getvalue())
    response.headers['Content-Type'] = 'application/pdf'
    response.headers['Content-Disposition'] = f'attachment; filename=financial_plan_{datetime.now().strftime("%Y%m%d")}.pdf'
    
    return response


@app.route('/api/export/<session_id>/docx', methods=['GET'])
def export_plan_docx(session_id):
    """Export planning summary as enhanced DOCX with professional formatting"""
    if session_id not in planning_sessions:
        return jsonify({'error': 'Session not found'}), 404
    
    session_data = planning_sessions[session_id]
    
    # Create DOCX document
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Title with enhanced styling
    title = doc.add_heading('Financial Planning Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.color.rgb = RGBColor(26, 54, 93)  # Dark blue
    
    # Subtitle
    date_para = doc.add_paragraph()
    date_run = date_para.add_run(f"Generated on {datetime.now().strftime('%B %d, %Y')}")
    date_run.font.size = Pt(12)
    date_run.font.color.rgb = RGBColor(74, 85, 104)  # Gray
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add spacing
    doc.add_paragraph()
    
    # Client Information Section
    client_heading = doc.add_heading('Client Information', level=1)
    client_heading.runs[0].font.color.rgb = RGBColor(45, 55, 72)  # Dark gray
    
    user_info = session_data['user_info']
    
    # Enhanced user information table
    user_table = doc.add_table(rows=4, cols=2)
    user_table.style = 'Light Grid Accent 1'
    
    # Header row styling
    header_cells = user_table.rows[0].cells
    header_cells[0].text = 'Client Details'
    header_cells[1].text = 'Information'
    
    # Make header bold
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(26, 54, 93)
    
    # Data rows
    data_rows = [
        ['Age', str(user_info.get('age', 'N/A'))],
        ['Annual Income', f"${user_info.get('annual_income', 0):,.2f}"],
        ['Current Savings', f"${user_info.get('savings', 0):,.2f}"]
    ]
    
    for i, (label, value) in enumerate(data_rows, 1):
        row_cells = user_table.rows[i].cells
        row_cells[0].text = label
        row_cells[1].text = value
        
        # Style label column
        label_run = row_cells[0].paragraphs[0].runs[0]
        label_run.font.bold = True
        
        # Style value column for financial data
        if '$' in value:
            value_run = row_cells[1].paragraphs[0].runs[0]
            value_run.font.color.rgb = RGBColor(49, 130, 206)  # Blue
    
    doc.add_paragraph()
    
    # Goals section
    if session_data.get('selected_plans'):
        goals_heading = doc.add_heading('Selected Financial Goals', level=2)
        goals_heading.runs[0].font.color.rgb = RGBColor(49, 130, 206)
        
        goals_para = doc.add_paragraph()
        for i, goal in enumerate(session_data['selected_plans']):
            if i > 0:
                goals_para.add_run(', ')
            goal_run = goals_para.add_run(goal)
            goal_run.font.italic = True
    
    doc.add_paragraph()
    
    # Executive Summary with enhanced formatting
    if 'Executive Summary' in session_data['plan_summaries']:
        exec_heading = doc.add_heading('Executive Summary', level=1)
        exec_heading.runs[0].font.color.rgb = RGBColor(45, 55, 72)
        
        exec_summary = clean_text_for_export(session_data['plan_summaries']['Executive Summary'])
        
        # Format executive summary with highlight box
        summary_para = doc.add_paragraph()
        summary_para.paragraph_format.space_before = Pt(6)
        summary_para.paragraph_format.space_after = Pt(12)
        
        # Add the summary text
        summary_run = summary_para.add_run(exec_summary)
        summary_run.font.size = Pt(11)
        
        doc.add_page_break()
    
    # Individual plans with enhanced formatting
    for plan_name, summary in session_data['plan_summaries'].items():
        if plan_name != 'Executive Summary':
            # Plan heading
            plan_heading = doc.add_heading(plan_name, level=1)
            plan_heading.runs[0].font.color.rgb = RGBColor(45, 55, 72)
            
            # Clean and format the summary
            clean_summary = clean_text_for_export(summary)
            
            # Parse content for better formatting
            lines = clean_summary.split('\n')
            current_para = None
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                # Check for bullet points
                if line.startswith(('•', '-', '*')):
                    bullet_para = doc.add_paragraph(line[1:].strip(), style='List Bullet')
                    bullet_para.paragraph_format.space_before = Pt(3)
                    bullet_para.paragraph_format.space_after = Pt(3)
                
                # Check for important information (contains financial data)
                elif any(char in line for char in ['$', '%']) or 'important' in line.lower():
                    important_para = doc.add_paragraph()
                    important_run = important_para.add_run(line)
                    important_run.font.bold = True
                    
                    # Highlight financial data
                    if any(char in line for char in ['$', '%']):
                        important_run.font.color.rgb = RGBColor(49, 130, 206)
                
                # Check for headings within the plan
                elif line.endswith(':') and len(line.split()) <= 5:
                    subheading_para = doc.add_paragraph()
                    subheading_run = subheading_para.add_run(line)
                    subheading_run.font.bold = True
                    subheading_run.font.size = Pt(12)
                    subheading_run.font.color.rgb = RGBColor(49, 130, 206)
                    subheading_para.paragraph_format.space_before = Pt(12)
                    subheading_para.paragraph_format.space_after = Pt(6)
                
                # Regular content
                else:
                    if current_para is None:
                        current_para = doc.add_paragraph()
                        current_para.paragraph_format.space_before = Pt(6)
                        current_para.paragraph_format.space_after = Pt(6)
                    
                    if current_para.text:
                        current_para.add_run(' ')
                    current_para.add_run(line)
            
            # Add spacing between plans
            doc.add_paragraph()
    
    # Footer section
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_run = footer_para.add_run(
        "Disclaimer: This financial plan is generated by AI and should be reviewed with a qualified financial advisor. "
        "For questions or updates, please consult your financial planning professional."
    )
    footer_run.font.size = Pt(9)
    footer_run.font.italic = True
    footer_run.font.color.rgb = RGBColor(113, 128, 150)
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save to buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    # Create response
    response = make_response(buffer.getvalue())
    response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    response.headers['Content-Disposition'] = f'attachment; filename=financial_plan_{datetime.now().strftime("%Y%m%d")}.docx'
    
    return response


# ============================================================================
# ERROR HANDLERS
# ============================================================================

@app.errorhandler(404)
def not_found(error):
    return jsonify({'error': 'Not found'}), 404


@app.errorhandler(500)
def server_error(error):
    return jsonify({'error': 'Internal server error'}), 500


# ============================================================================
# MAIN
# ============================================================================

if __name__ == '__main__':
    # Verify API key is set
    if not app.config['OPENAI_API_KEY']:
        raise ValueError("OPENAI_API_KEY not set in environment variables. Please create a .env file with your API key.")
    
    port = app.config['APP_PORT']
    host = app.config['APP_HOST']
    
    print(f"\n{'='*60}")
    print(f"Financial Planner Web Application")
    print(f"{'='*60}")
    print(f"Starting server on http://{host}:{port}")
    print(f"{'='*60}\n")
    
    app.run(host=host, port=port, debug=app.config['DEBUG'])
